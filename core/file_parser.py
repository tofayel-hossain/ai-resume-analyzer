from io import BytesIO
from pathlib import Path

import pymupdf
from docx import Document


class ResumeParseError(ValueError):
    pass


def extract_pdf_text(file_bytes: bytes) -> str:
    try:
        doc = pymupdf.open(stream=file_bytes, filetype="pdf")
        parts = []
        for page in doc:
            parts.append(page.get_text("text", sort=True))
        doc.close()
        text = "\n".join(parts).strip()
    except Exception as exc:
        raise ResumeParseError(f"Could not read PDF: {exc}") from exc

    if not text:
        raise ResumeParseError(
            "No selectable text was found in the PDF. "
            "This MVP does not OCR image-only/scanned resumes."
        )
    return text


def extract_docx_text(file_bytes: bytes) -> str:
    try:
        doc = Document(BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        text = "\n".join(parts).strip()
    except Exception as exc:
        raise ResumeParseError(f"Could not read DOCX: {exc}") from exc

    if not text:
        raise ResumeParseError("The DOCX did not contain readable text.")
    return text


def extract_resume_text(file_bytes: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_pdf_text(file_bytes)
    if ext == ".docx":
        return extract_docx_text(file_bytes)
    raise ResumeParseError("Only PDF and DOCX files are supported.")
